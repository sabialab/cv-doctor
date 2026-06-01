"""DOCX exporter tests."""

from __future__ import annotations

import uuid
from pathlib import Path

from docx import Document

from src.models import Change, ChangeRisk, ChangeStatus
from src.services.exporter_docx import apply_changes_to_docx


def test_apply_changes_to_docx(tmp_path: Path):
    from io import BytesIO

    doc = Document()
    doc.add_paragraph("熟悉 Python")
    buf = BytesIO()
    doc.save(buf)
    resume_bytes = buf.getvalue()

    out = tmp_path / "out.docx"
    changes = [
        Change(
            id=str(uuid.uuid4()),
            section="skills",
            original="熟悉 Python",
            revised="Python（3 年）",
            reason="test",
            evidence_ids=["f1"],
            risk_level=ChangeRisk.LOW,
            status=ChangeStatus.ACCEPTED,
        )
    ]
    applied = apply_changes_to_docx(resume_bytes, changes, out)
    assert applied == 1
    assert out.is_file()

    result = Document(out)
    assert "Python（3 年）" in result.paragraphs[0].text
