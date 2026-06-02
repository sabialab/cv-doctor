"""P0 API 冒烟测试（桩流水线）。"""

from fastapi.testclient import TestClient

from src.main import app


def test_create_session_requires_consent():
    client = TestClient(app)
    from io import BytesIO

    from docx import Document

    buf = BytesIO()
    doc = Document()
    doc.add_paragraph("test")
    doc.save(buf)
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    files = {"resume": ("resume.docx", buf.getvalue(), mime)}
    r = client.post(
        "/sessions",
        files=files,
        data={"jd_text": "需要 Python", "consent": "false"},
    )
    assert r.status_code == 400
    assert "隐私" in r.json()["detail"]


def test_health():
    client = TestClient(app)
    r = client.get("/health")
    assert r.status_code == 200
    assert r.json()["status"] == "ok"


def test_session_flow_stub():
    client = TestClient(app)
    # 使用 python-docx 生成最小 docx
    from io import BytesIO

    from docx import Document

    buf = BytesIO()
    doc = Document()
    doc.add_paragraph("负责后端开发与维护。")
    doc.add_paragraph("熟悉 Python")
    doc.save(buf)
    buf.seek(0)
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    files = {"resume": ("resume.docx", buf.read(), mime)}
    data = {
        "jd_text": "需要 Python 和 FastAPI 经验的后端工程师。",
        "consent": "true",
    }
    r = client.post("/sessions", files=files, data=data)
    assert r.status_code == 200, r.text
    sid = r.json()["session_id"]

    # 同步 TestClient 会跑完 background task
    g = client.get(f"/sessions/{sid}")
    assert g.status_code == 200
    body = g.json()
    assert body["status"] == "ready"
    assert body["result"] is not None
    assert len(body["result"]["changes"]) <= 3
    assert "policy_guard" in body["result"]
    assert body["result"]["policy_guard"]["passed"] is True
    assert body["result"].get("free_change_limit") == 3

    cid = body["result"]["changes"][0]["id"]
    p = client.patch(f"/sessions/{sid}/changes/{cid}", json={"status": "accepted"})
    assert p.status_code == 200

    e = client.post(f"/sessions/{sid}/export")
    assert e.status_code == 200
    assert e.json()["format"] == "docx"
    dl = e.json()["download_url"]
    down = client.get(dl)
    assert down.status_code == 200
    assert "wordprocessingml" in down.headers.get("content-type", "")

    d = client.delete(f"/sessions/{sid}")
    assert d.status_code == 200
