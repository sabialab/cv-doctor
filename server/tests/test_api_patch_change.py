"""PATCH /sessions/{id}/changes/{id} — status and revised text."""

from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from src.main import app


def _minimal_docx() -> bytes:
    buf = BytesIO()
    doc = Document()
    doc.add_paragraph("负责后端开发与维护。")
    doc.add_paragraph("熟悉 Python")
    doc.save(buf)
    return buf.getvalue()


def _ready_session(client: TestClient) -> tuple[str, str, str]:
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    files = {"resume": ("resume.docx", _minimal_docx(), mime)}
    data = {"jd_text": "需要 Python 和 FastAPI 经验的后端工程师。"}
    sid = client.post("/sessions", files=files, data=data).json()["session_id"]
    body = client.get(f"/sessions/{sid}").json()
    ch0 = body["result"]["changes"][0]
    return sid, ch0["id"], ch0["original"]


def _ready_medium_risk_change(client: TestClient) -> tuple[str, str, str]:
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    files = {"resume": ("resume.docx", _minimal_docx(), mime)}
    data = {"jd_text": "需要 Python 和 FastAPI 经验的后端工程师。"}
    sid = client.post("/sessions", files=files, data=data).json()["session_id"]
    body = client.get(f"/sessions/{sid}").json()
    medium = next(c for c in body["result"]["changes"] if c["risk_level"] == "medium")
    return sid, medium["id"], medium["original"]


def test_patch_revised_only_updates_export_docx():
    client = TestClient(app)
    sid, cid, original = _ready_session(client)
    edited = f"{original.rstrip('。')}（用户编辑稿）。"
    r = client.patch(f"/sessions/{sid}/changes/{cid}", json={"revised": edited})
    assert r.status_code == 200
    got = client.get(f"/sessions/{sid}").json()
    ch = next(c for c in got["result"]["changes"] if c["id"] == cid)
    assert ch["revised"] == edited
    assert ch["status"] == "accepted"
    exp = client.post(f"/sessions/{sid}/export")
    assert exp.status_code == 200
    down = client.get(exp.json()["download_url"])
    assert down.status_code == 200
    out = Document(BytesIO(down.content))
    blob = "\n".join(p.text for p in out.paragraphs)
    assert "用户编辑稿" in blob


def test_patch_medium_risk_revised_requires_separate_accept():
    client = TestClient(app)
    sid, cid, original = _ready_medium_risk_change(client)
    edited = f"{original.rstrip('。')}（用户编辑稿）。"

    r = client.patch(f"/sessions/{sid}/changes/{cid}", json={"revised": edited})

    assert r.status_code == 200
    got = client.get(f"/sessions/{sid}").json()
    ch = next(c for c in got["result"]["changes"] if c["id"] == cid)
    assert ch["revised"] == edited
    assert ch["requires_user_confirmation"] is True
    assert ch["status"] == "pending"
    exp = client.post(f"/sessions/{sid}/export")
    assert exp.status_code == 400
    assert "请先接受" in exp.json()["detail"]


def test_patch_rejects_revised_and_status_together():
    client = TestClient(app)
    sid, cid, _ = _ready_session(client)
    r = client.patch(
        f"/sessions/{sid}/changes/{cid}",
        json={"revised": "x", "status": "rejected"},
    )
    assert r.status_code == 422


def test_patch_rejects_forbidden_revised():
    client = TestClient(app)
    sid, cid, original = _ready_session(client)
    r = client.patch(
        f"/sessions/{sid}/changes/{cid}",
        json={"revised": f"{original}（编造经历）"},
    )
    assert r.status_code == 400
    ch = next(
        c for c in client.get(f"/sessions/{sid}").json()["result"]["changes"] if c["id"] == cid
    )
    assert ch["status"] == "pending"


def test_patch_rejects_whitespace_only_revised():
    client = TestClient(app)
    sid, cid, _ = _ready_session(client)
    r = client.patch(f"/sessions/{sid}/changes/{cid}", json={"revised": "   "})
    assert r.status_code == 422


def test_patch_change_clears_stale_export():
    client = TestClient(app)
    sid, cid, _ = _ready_session(client)
    client.patch(f"/sessions/{sid}/changes/{cid}", json={"status": "accepted"})
    exp = client.post(f"/sessions/{sid}/export")
    assert exp.status_code == 200
    client.patch(f"/sessions/{sid}/changes/{cid}", json={"status": "rejected"})
    down = client.get(exp.json()["download_url"])
    assert down.status_code == 404


def test_patch_status_only_reject():
    client = TestClient(app)
    sid, cid, _ = _ready_session(client)
    r = client.patch(f"/sessions/{sid}/changes/{cid}", json={"status": "rejected"})
    assert r.status_code == 200
    ch = next(
        c for c in client.get(f"/sessions/{sid}").json()["result"]["changes"] if c["id"] == cid
    )
    assert ch["status"] == "rejected"
