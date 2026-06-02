"""Apply accepted changes to original resume DOCX."""

from __future__ import annotations

from collections.abc import Iterator
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

from src.models import Change


def _iter_paragraphs(doc: Document) -> Iterator[Paragraph]:
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


def _replace_in_paragraph(para: Paragraph, original: str, revised: str) -> bool:
    if original not in para.text:
        return False
    new_text = para.text.replace(original, revised, 1)
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.text = new_text
    return True


def build_docx_from_plain_text(text: str) -> bytes:
    """Build a minimal DOCX from pasted resume text (text-only sessions)."""
    buf = BytesIO()
    doc = Document()
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        doc.add_paragraph(text.strip() or "")
    else:
        for line in lines:
            doc.add_paragraph(line)
    doc.save(buf)
    return buf.getvalue()


def apply_changes_to_docx(
    resume_bytes: bytes,
    changes: list[Change],
    output_path: Path,
) -> int:
    """Replace `original` with `revised` in body/table paragraphs. Returns applied count."""
    doc = Document(BytesIO(resume_bytes))
    applied = 0
    for change in changes:
        if not change.original.strip():
            continue
        for para in _iter_paragraphs(doc):
            if _replace_in_paragraph(para, change.original, change.revised):
                applied += 1
                break
    doc.save(output_path)
    return applied
